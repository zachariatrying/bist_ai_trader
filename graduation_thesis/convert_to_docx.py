import os
import sys
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

thesis_dir = os.path.dirname(os.path.abspath(__file__)).replace("\\", "/")
md_path = os.path.join(thesis_dir, "Tez_Metni_Final.md")
docx_path = os.path.join(thesis_dir, "BorsaNeuron_Graduation_Thesis.docx")

def set_cell_background(cell, fill_hex):
    """Set cell background shading."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding/margins for table cells in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def parse_markdown_text(paragraph, text):
    """Parse basic markdown bold (**text**) and italic (*text*) and convert to docx runs."""
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\)|`.*?`|[^\*`\[]+)')
    matches = pattern.findall(text)
    
    for match in matches:
        if match.startswith('***') and match.endswith('***'):
            run = paragraph.add_run(match[3:-3])
            run.bold = True
            run.italic = True
        elif match.startswith('**') and match.endswith('**'):
            run = paragraph.add_run(match[2:-2])
            run.bold = True
        elif match.startswith('*') and match.endswith('*'):
            run = paragraph.add_run(match[1:-1])
            run.italic = True
        elif match.startswith('`') and match.endswith('`'):
            run = paragraph.add_run(match[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10.5)
        elif match.startswith('[') and ']' in match:
            link_text = match[1:match.find(']')]
            run = paragraph.add_run(link_text)
            run.font.color.rgb = RGBColor(0, 0, 238)
            run.underline = True
        else:
            paragraph.add_run(match)

def apply_global_styles(doc):
    """Apply Yeditepe University margins and styles to document."""
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Configure footer
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f_run = f_p.add_run("BorsaNeuron Graduation Thesis  |  Page ")
        f_run.font.name = 'Times New Roman'
        f_run.font.size = Pt(9.5)
        f_run.font.color.rgb = RGBColor(120, 120, 120)
        
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'PAGE')
        f_p._p.append(fldSimple)

def add_p_shading(p, fill_hex):
    """Add paragraph background shading (for code blocks)."""
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    pPr.append(shd)

def build_docx():
    print(f"Reading markdown from: {md_path}")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    doc = docx.Document()
    apply_global_styles(doc)
    
    is_cover_page = True
    in_table = False
    in_code_block = False
    is_first_h1 = True
    last_was_blank = False
    table_lines = []
    
    for index, line in enumerate(lines):
        clean_line = line.strip()
        
        # 1. Handle horizontal rule for cover page separation
        if clean_line == "---" and is_cover_page:
            is_cover_page = False
            doc.add_page_break()
            continue
            
        if is_cover_page:
            is_br = re.match(r'^<br\s*/?>+$', clean_line.replace(" ", "").lower()) is not None
            if not clean_line or is_br:
                if not last_was_blank:
                    doc.add_paragraph()
                    last_was_blank = True
                continue
            
            last_was_blank = False
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Remove HTML line breaks from headings/text if any
            display_text = re.sub(r'<br\s*/?>', '', clean_line, flags=re.IGNORECASE).strip()
            
            if display_text.startswith("# "):
                run = p.add_run(display_text[2:])
                run.bold = True
                run.font.size = Pt(20)
            elif display_text.startswith("## "):
                run = p.add_run(display_text[3:])
                run.bold = True
                run.font.size = Pt(16)
            elif display_text.startswith("### "):
                run = p.add_run(display_text[4:])
                run.bold = True
                run.font.size = Pt(14)
            else:
                run = p.add_run(display_text)
                run.font.size = Pt(12)
                
            run.font.name = 'Times New Roman'
            continue

        # 2. Parse Code Blocks
        if clean_line.startswith("```"):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            add_p_shading(p, "F4F4F5")
            
            run = p.add_run(line.replace("\n", ""))
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(60, 60, 60)
            continue

        # 3. Parse Tables
        if clean_line.startswith("|"):
            in_table = True
            table_lines.append(clean_line)
            continue
        elif in_table and not clean_line.startswith("|"):
            in_table = False
            if table_lines:
                processed_lines = [l for l in table_lines if not re.match(r'^\|[\s\-\|]+$', l)]
                rows_data = []
                for t_line in processed_lines:
                    cols = [c.strip() for c in t_line.split("|")[1:-1]]
                    rows_data.append(cols)
                
                if rows_data:
                    num_cols = len(rows_data[0])
                    num_rows = len(rows_data)
                    
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.autofit = True
                    
                    for r_idx, row_cells in enumerate(table.rows):
                        row_data = rows_data[r_idx]
                        for c_idx, cell in enumerate(row_cells.cells):
                            if c_idx < len(row_data):
                                cell.text = ""
                                p = cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                run = p.add_run(row_data[c_idx])
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(10.5)
                                
                                if r_idx == 0:
                                    run.bold = True
                                    set_cell_background(cell, "F2F2F2")
                                
                                set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
            
            table_lines = []
            if not clean_line:
                continue

        # 4. Headings & Spacings
        if clean_line.startswith("# "):
            if not is_first_h1:
                doc.add_page_break()
            else:
                is_first_h1 = False
                
            p = doc.add_heading(level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            
            run = p.add_run(clean_line[2:])
            run.font.name = 'Times New Roman'
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(17, 17, 17)
            continue
            
        elif clean_line.startswith("## "):
            p = doc.add_heading(level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
            run = p.add_run(clean_line[3:])
            run.font.name = 'Times New Roman'
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(30, 30, 30)
            continue
            
        elif clean_line.startswith("### "):
            p = doc.add_heading(level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            
            run = p.add_run(clean_line[4:])
            run.font.name = 'Times New Roman'
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(50, 50, 50)
            continue

        # 5. Handle embedded images
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)$', clean_line)
        if img_match:
            alt_text = img_match.group(1)
            img_rel_path = img_match.group(2)
            img_abs_path = os.path.join(thesis_dir, img_rel_path)
            
            if os.path.exists(img_abs_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_abs_path, width=Inches(6.0))
                
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap_run = p_cap.add_run(f"Figure: {alt_text}")
                p_cap_run.italic = True
                p_cap_run.font.name = 'Times New Roman'
                p_cap_run.font.size = Pt(10)
                p_cap_run.font.color.rgb = RGBColor(80, 80, 80)
            else:
                print(f"Warning: Image file not found at: {img_abs_path}")
            continue

        # 6. List items
        list_match = re.match(r'^[\-\*\+]\s+(.*)$', clean_line)
        if list_match:
            content = list_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            parse_markdown_text(p, content)
            
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            continue

        # 7. Blockquotes
        quote_match = re.match(r'^>\s*(.*)$', clean_line)
        if quote_match:
            content = quote_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(6)
            parse_markdown_text(p, content)
            
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(80, 80, 80)
                run.italic = True
            continue

        # 8. Standard Paragraphs
        if clean_line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(8)
            parse_markdown_text(p, clean_line)
            
            for run in p.runs:
                if run.font.name != 'Courier New':
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

    print(f"Saving compiled MS Word document to: {docx_path}")
    success = False
    try:
        doc.save(docx_path)
        print("MS Word document successfully generated!")
        success = True
    except PermissionError:
        print(f"Permission denied on {docx_path}.")
        
    if not success:
        i = 2
        while True:
            alt_docx_path = docx_path.replace(".docx", f"_v{i}.docx")
            try:
                doc.save(alt_docx_path)
                print(f"Alternative MS Word document successfully generated at: {alt_docx_path}")
                break
            except PermissionError:
                print(f"Permission denied on {alt_docx_path}.")
                i += 1

if __name__ == "__main__":
    build_docx()
